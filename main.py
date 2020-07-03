
#################################################################################################################################################
from docx.shared import RGBColor
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import RGBColor
import urllib.request
from urllib.request import *
from docx.oxml.ns import qn

#################################################################################################################################################
def formatting_novels(novel_directory_and_name):  
    with open(novel_directory_and_name,'r') as f:
        a = f.readlines()
    a =  list(''.join(a))
    newlist = []
    #throw = []
    i = 0
    while i<=len(a)-1 and i+8<=len(a)-1:
        if ord(a[i])<=180 and ord(a[i+1])<=180 and ord(a[i+2])<=180 and ord(a[i+3])<=180 and ord(a[i+4])<=180 and ord(a[i+5])<=180  and ord(a[i+6])<=180  and ord(a[i+7])<=180  and ord(a[i+8])<=180:
            newlist.append('0')
        else:
            newlist.append(a[i])
        i = i+1
    q = (''.join(newlist))
    num = 1020
    while num>=15:
        q = q.replace(num*'0','\n')
        q = q.replace('/div&gt;','')
        q = q.replace('p;nbsp;','')
        q = q.replace('mp;nbsp;','')
        q = q.replace('p;nbsp;','')
        q = q.replace('uot;&gt;','')
        q = q.replace('t;/a&gt;','')
        q = q.replace('&quot;','"')
        q = q.replace('「','“')
        q = q.replace('」','”')
        q = q.replace('『','“')
        q = q.replace('』','”')
        num = num-1
    return (q)
################################################################################################################################################
def get_link(a,num):
    link = 'http://'+a+'.blogspot.com/feeds/posts/default?start-index='+str(num)
    result = urlopen(link)
    print ('opening link successfully')
    data = result.read()
    with open(str(num)+'.txt','w',errors='ignore') as f:
        f.write(data.decode('utf-8'))
        print ('writing to designated folder successfully')
################################################################################################################################################
def breakdown_novel(file): ###拆分文档
    with open(file,'r',encoding = 'gbk',errors ='ignore') as f:
        a = f.readlines()
    a = ''.join(a)
    key = 1 #### to record the number of splitted novels
    while key<=99:
        try:
            start = "<title type='text'>"
            ba = start + a[a.index(start)+len(start):a.index(start)+len(start)+15]
            queer = '*'*100 + a[a.index(start)+len(start):a.index(start)+len(start)+15]
            a = a.replace(ba,queer)
            end =  "<title type='text'>"
            with open(str(key)+'.txt','w') as f:
                spec_novel = ((a[a.index('************'):a.index(end)]))
                spec_novel = spec_novel.replace('************','')
                f.write(spec_novel)
            spec_novel2 = ((a[a.index('************'):a.index(end)-len(end)]))
            a = a.replace(spec_novel2,'')
            key = key+1
        except:
            break
    print ('breaking down novel successfully')
################################################################################################################################################    
import os
def file_name(file_dir,types):
    L=[]
    for root, dirs, files in os.walk(file_dir):
        for file in files:
            if os.path.splitext(file)[1] == types:# '.txt':
                L.append(os.path.join(root, file))
    return L
################################################################################################################################################
def get_title(novel_directory_and_name):
    with open(novel_directory_and_name,encoding = 'gbk',errors = 'ignore') as f:
        a = f.readlines()
    a = ''.join(a)
    start = '****'                            ###
    end = '</title><content type'             ### 
    title = a[a.index(start)+4:a.index(end)]  ###
    qr ="</title><content type='html'>&lt;div class=&quot;separator&quot; style=&quot;clear: both; text-align: center;&quot;&gt;&lt;a href=&quot;"
    a = a.replace(qr,'')
    return(title)
################################################################################################################################################
def convert_to_docx(title,content):
    document = Document()
    p = document.add_paragraph('')
    rush = p.add_run(title)
    font = rush.font
    font.size = Pt(20) # 
    font.color.rgb = RGBColor(0,0,0)
    rust = p.add_run(content)
    font = rush.font
    font.size = Pt(15) # 
    font.color.rgb = RGBColor(0,0,0)
    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    document.save(title+'.docx')
    print ('Your converting process has been completed!')
################################################################################################################################################
def title_cleaning(a):
    blacklist = list('\/:*?"<> |')
    for i in blacklist:
        a = a.replace(i,'')
    return (a)

################################################################################################################################################
from langconv import Converter
def convert(text,flag=0):
    rule = 'zh-hans' if flag else 'zh-hans'
    return Converter(rule).convert(text)
################################################################################################################################################
import time
from os.path import getsize
import os
import shutil

loc = input('please indicate a folder where you would save your blog, I will recommend you start in a space where no .txt is found, otherwise all .txt will be removed when the program starts!')
temp = input('indicate a folder where you would save a log.txt file onto your computer, you have to remember the location!  You will be prompted to enter the number of log when you open this program again!')
log = input('what is your log? If you do not know, please enter 1,always make sure you enter an integer')
num = int(log)
while num<=9999:
    afe = time.time()
    a = input("Please enter the blogspot title,for instance, if it is text.blogspot.com, then you should write text, I'll download every article in this blogspot then.")
    get_link(a,num)
    breakdown_novel(str(num)+'.txt')
    g = file_name(loc,'.txt')
    g.remove(loc+'\\'+str(num)+'.txt')
    for i in g:
        try:
            title = get_title(i)
            title = title_cleaning(title)
            title = convert(title)
            if (loc+'\\'+title+'.docx' in file_name(loc,'.docx')) ==True:
                fuli = 0
            else:
                content = formatting_novels(i)
                content = convert(content)
                convert_to_docx(title,content)
                print ('convert success')
        except:
            if getsize(i)==0:
                print ('this file is empty, so I will not convert it ')
            else:
                with open(i,'r') as f:
                    a = f.readlines()
                a = ''.join(a)
                print (' I am directing your file to another location for future checkups！')
                shutil.copy(i,temp)
    g = file_name(loc,'.txt')
    for i in g:
        os.remove(i)
    with open (temp+'\\'+log.txt','w')as f:
        f.write(str(num))
    bfe = time.time()
    print ('using time of seconds:',int(bfe-afe))
    num = num+1
    print ('Your log is....',str(num),"when you exit the process, remember this word on a piece of paper or so, so that when you want to download the same blog under the same folder later, you could enter the log number to continue")
